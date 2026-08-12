# infrastructure/adapters/docx_generator.py
import io  # <--- Agregar esta importación
import re  # <--- Import regex module
from docxtpl import DocxTemplate, RichText
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from application.ports.outbound import DocumentGeneratorPort
from domain.models import SessionContext, TopicDetail, SubtopicRow
from typing import List, Union
import copy
import roman

class DocxTemplateGeneratorAdapter(DocumentGeneratorPort):
    def generate(self, context: SessionContext, template_path: str) -> bytes:
        
        # --- PASO 1: Renderizar datos fijos ---
        doc_tpl = DocxTemplate(template_path)
        scalar_payload = {
            "colegio": context.datos_generales.colegio,
            "area": context.datos_generales.area,
            "nivel": context.datos_generales.nivel,
            "grado": context.datos_generales.grado,
            "seccion": context.datos_generales.seccion, 
            "docente": context.datos_generales.docente,
            "mes": context.datos_generales.mes,
            "valor": context.valor.upper(),
            "topics": self.getTopics(context.temas)
        }

        doc_tpl.render(scalar_payload)

        # Guardamos la primera pasada en un buffer en memoria
        intermediate_buffer = io.BytesIO()
        doc_tpl.save(intermediate_buffer)
        intermediate_buffer.seek(0)        

        # --- PASO 2: Rellenar las tablas manteniendo su estilo uniforme ---
        doc = Document(intermediate_buffer)

        for table in doc.tables:
            if len(table.rows) >= 2:
                first_cell = table.rows[1].cells[0]
                cell_text = "".join(p.text for p in first_cell.paragraphs).strip()
                
                if "[TABLA_PROPOSITOS]" in cell_text:
                    # Guardamos la fila 1 (donde está el marcador) como nuestro molde definitivo
                    base_row = table.rows[1]
                    
                    # Iteramos sobre TODOS los datos y les creamos una fila NUEVA a cada uno
                    for p in context.proposito:
                        new_row = table.add_row()
                        
                        # Copiamos los estilos XML internos de celda a celda (colores, bordes)
                        for i, cell in enumerate(new_row.cells):
                            cell._tc.get_or_add_tcPr().append(
                                copy.deepcopy(base_row.cells[i]._tc.get_or_add_tcPr())
                            )
                        
                        # Limpiamos párrafos residuales en la nueva fila y dejamos solo uno
                        for c in new_row.cells:
                            if not c.paragraphs:
                                c.add_paragraph()
                            c.paragraphs[0].text = ""
                            while len(c.paragraphs) > 1:
                                p_element = c.paragraphs[-1]._p
                                p_element.getparent().remove(p_element)

                        # Escribimos los datos en la fila nueva
                        self._set_cell_text(new_row.cells[0], str(p.competencia))
                        # Aplicamos viñetas a Capacidades y Desempeños
                        self._set_cell_bullets(new_row.cells[1], p.capacidades)
                        self._set_cell_bullets(new_row.cells[2], p.desempenos)

                    
                    # --- EL TRUCO FINAL ---
                    # Ahora que ya creamos todas las filas con datos reales usando el molde...
                    # ¡Eliminamos la fila base original que contenía el problema de altura!
                    tr_element = base_row._tr
                    tr_element.getparent().remove(tr_element)

                if "[TABLA_ENFOQUES]" in cell_text:
                    base_row = table.rows[1]
                    current_tr = base_row._tr  # Puntero para mantener la posición exacta en el XML
                    
                    for p in context.enfoques:
                        new_row = table.add_row()
                        
                        # Mover la fila recién creada a la posición inmediatamente superior a VALOR
                        current_tr.addnext(new_row._tr)
                        current_tr = new_row._tr  # Avanzar la posición del puntero
                        
                        for i, cell in enumerate(new_row.cells):
                            cell._tc.get_or_add_tcPr().append(
                                copy.deepcopy(base_row.cells[i]._tc.get_or_add_tcPr())
                            )
                        
                        for c in new_row.cells:
                            if not c.paragraphs:
                                c.add_paragraph()
                            c.paragraphs[0].text = ""
                            while len(c.paragraphs) > 1:
                                p_element = c.paragraphs[-1]._p
                                p_element.getparent().remove(p_element)
                        # Set cell text with Arial 12pt
                        self._set_cell_text(new_row.cells[0], str(p.enfoque))
                        self._set_cell_text(new_row.cells[1], str(p.acciones))
                    
                    # --- EL TRUCO FINAL ---
                    # Ahora que ya creamos todas las filas con datos reales usando el molde...
                    # ¡Eliminamos la fila base original que contenía el problema de altura!
                    tr_element = base_row._tr
                    tr_element.getparent().remove(tr_element)

        # --- PASO 3: Guardar el resultado en el buffer final en memoria ---
        final_buffer = io.BytesIO()
        doc.save(final_buffer)
        final_buffer.seek(0)

        print("buffer", final_buffer)
        
        return final_buffer

    # --- HELPER TO PARSE MARKDOWN BOLD (**text**) ---
    def _append_formatted_text(self, paragraph, text: str) -> None:
        """Splits markdown **bold** text and applies run.bold = True where necessary."""
        # Split text by **...** tags
        parts = re.split(r'\*\*(.*?)\*\*', text)
        
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            # Odd index positions (1, 3, 5...) correspond to the captured content inside **
            if idx % 2 == 1:
                run.bold = True    

    def _set_cell_text(self, cell, text: str) -> None:
        """Helper to write plain text to a cell with explicit Arial 12pt formatting."""
        p = cell.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Explicitly set left alignment
        self._append_formatted_text(p, text)

    def _set_cell_bullets(self, cell, items: Union[List[str], str]) -> None:
        """Helper to format bullet points inside table cells with explicit Arial 12pt."""
        if isinstance(items, str):
            item_list = [line.strip() for line in items.split("\n") if line.strip()]
        elif isinstance(items, list):
            item_list = [str(i).strip() for i in items if str(i).strip()]
        else:
            item_list = [str(items)] if items else []

        if not item_list:
            return

        # Configure the first paragraph
        p0 = cell.paragraphs[0]
        p0.text = ""
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p0.style = 'List Bullet'
            self._append_formatted_text(p0, item_list[0])
        except KeyError:
            p0.add_run("• ")
            self._append_formatted_text(p0, item_list[0])

        # Add remaining paragraphs
        for item in item_list[1:]:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p.style = 'List Bullet'
                run = p.add_run(item)
            except KeyError:
                run = p.add_run(f"• {item}")
            
            run.font.name = 'Arial'
            run.font.size = Pt(12)

    # --- HELPER TO PARSE MARKDOWN BOLD FOR PYTHON-DOCX PARAGRAPHS ---
    def _append_formatted_text(self, paragraph, text: str) -> None:
        """Splits markdown **bold** text and applies run.bold = True where necessary."""
        parts = re.split(r'\*\*(.*?)\*\*', str(text))
        
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            if idx % 2 == 1:
                run.bold = True

    # --- HELPER TO CONVERT MARKDOWN BOLD TO DOCXTPL RICHTEXT ---
    def _to_rich_text(self, text: str) -> RichText:
        """Converts Markdown **bold** text into a docxtpl RichText object."""
        if not text:
            return RichText("")

        rt = RichText()
        parts = re.split(r'\*\*(.*?)\*\*', str(text))

        for idx, part in enumerate(parts):
                    if not part:
                        continue
                    
                    # Odd indexes correspond to captured content inside **
                    is_bold = (idx % 2 == 1)
                    
                    # size=24 corresponds to 12pt in Word XML half-points
                    rt.add(part, font='Arial', size=24, bold=is_bold)

        return rt            

    def getTopics(self, topics: List[TopicDetail]) -> List:
        temas = []
        for index, topic in enumerate(topics):
                letra = chr(65 + index)  # A, B, C...
                # Safe extraction of attributes
                # index + 1 porque los números romanos empiezan en 1 (I, II, III...)
                # .lower() convierte el resultado a minúsculas (i, ii, iii...)
                letra_romana = roman.toRoman(index + 1).upper()                
                
                temas.append({
                    "letra": letra,
                    "roman": letra_romana,
                    "titulo": topic.titulo,
                    "inicio": self._to_rich_text(topic.inicio),
                    "cierre": topic.cierre,
                    "tarea": self._to_rich_text(topic.tarea),
                    "subtemas": self.getSubTopics(topic.subtemas)
                })
        return temas
    
    def getSubTopics(self, subTopics: List[SubtopicRow]) -> List:
        subTemas = []
        for index, subTopic in enumerate(subTopics):
            
            subTemas.append({
                "index": index + 1,
                "subtema": subTopic.subtema,
                "resumen": self._to_rich_text(subTopic.resumen)
            })
        return subTemas